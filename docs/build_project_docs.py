"""
Build PDF and Word documents from docs/PROJECT_EXPLANATION.md
Run from project root: python docs/build_project_docs.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "PROJECT_EXPLANATION.md"


def _escape_reportlab(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_pdf(md_text: str, out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=9,
        leading=11,
        spaceAfter=6,
    )
    h1 = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=14,
    )
    h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=11,
        spaceAfter=8,
        spaceBefore=10,
    )
    mono = ParagraphStyle(
        "Mono",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=7,
        leading=9,
        leftIndent=0,
    )

    story = []
    blocks = re.split(r"\n{2,}", md_text)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        first = lines[0].strip()
        if first.startswith("# "):
            story.append(Paragraph(_escape_reportlab(first[2:].strip()), h1))
        elif first.startswith("## "):
            story.append(Paragraph(_escape_reportlab(first[3:].strip()), h2))
        elif first.startswith("```"):
            code = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
            for line in code.split("\n"):
                story.append(Paragraph(_escape_reportlab(line[:2000]), mono))
            story.append(Spacer(1, 0.15 * cm))
        else:
            para = " ".join(lines)
            story.append(Paragraph(_escape_reportlab(para), body))

    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    doc.build(story)


def build_docx(md_text: str, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    blocks = re.split(r"\n{2,}", md_text)
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        first = lines[0].strip()
        if first.startswith("# "):
            p = doc.add_heading(first[2:].strip(), level=1)
        elif first.startswith("## "):
            p = doc.add_heading(first[3:].strip(), level=2)
        elif first.startswith("### "):
            p = doc.add_heading(first[4:].strip(), level=3)
        elif first.startswith("```"):
            code = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
            p = doc.add_paragraph(code)
            for run in p.runs:
                run.font.name = "Courier New"
        else:
            doc.add_paragraph(" ".join(lines))

    doc.save(out_path)


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing {MD_PATH}", file=sys.stderr)
        return 1
    text = MD_PATH.read_text(encoding="utf-8")
    pdf_out = BASE / "PROJECT_EXPLANATION.pdf"
    docx_out = BASE / "PROJECT_EXPLANATION.docx"
    build_pdf(text, pdf_out)
    print(f"Wrote {pdf_out}")
    build_docx(text, docx_out)
    print(f"Wrote {docx_out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
